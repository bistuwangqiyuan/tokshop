"""把报告 Markdown 转成 Word (.docx)。

为什么不用 pandoc
-----------------
本机没有 pandoc，而装它需要网络与管理员权限。python-docx 已在环境里，且自建
转换器有两个实际好处：中文字体可以精确控制（东亚字体要单独设 w:eastAsia，
否则 Word 会用西文字体渲染中文导致字形难看），表格样式与图片缩放也能按报告
的实际需要来调，而不是接受 pandoc 的默认值。

支持的 Markdown 子集（正好覆盖本仓报告用到的全部语法）
------------------------------------------------------
  - ATX 标题 # ~ ####
  - 段落，行内 **粗体**、`代码`、[链接](url)
  - GFM 管道表格，含 :---: 对齐
  - 无序列表 -、有序列表 1.
  - 围栏代码块 ```lang，其中 ```mermaid 若存在同名预渲染 PNG 则改为插图
  - 图片 ![alt](相对路径)
  - 引用块 >
  - 分隔线 ---

用法
----
    python scripts/md2docx.py reports/sea-b300-kimi-k3/01-*.md
    python scripts/md2docx.py --all          # 转换 reports/ 下全部 .md
    python scripts/md2docx.py --all --outdir dist/docx
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# 正文中文用宋体、标题用黑体、代码用 Consolas 配等宽中文
FONT_BODY_CN = "宋体"
FONT_HEAD_CN = "微软雅黑"
FONT_BODY_EN = "Times New Roman"
FONT_HEAD_EN = "Segoe UI"
FONT_MONO = "Consolas"

ACCENT = RGBColor(0x1F, 0x3B, 0x63)      # 标题深蓝
MUTED = RGBColor(0x60, 0x60, 0x60)       # 注释灰
CODE_BG = "F2F3F5"
TABLE_HEAD_BG = "1F3B63"


# ---------------------------------------------------------------------------
# 字体与底纹的低层工具
# ---------------------------------------------------------------------------

def set_run_font(run, *, cn: str, en: str, size: float | None = None,
                 bold: bool | None = None, color: RGBColor | None = None,
                 mono: bool = False) -> None:
    """同时设置西文与东亚字体。

    python-docx 的 run.font.name 只写 w:ascii/w:hAnsi，中文会落到默认字体。
    必须显式写 w:eastAsia，否则中文在 Word 里会被渲染成不匹配的字形。
    """
    if mono:
        cn = en = FONT_MONO
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    rfonts.set(qn("w:eastAsia"), cn)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(element, fill: str) -> None:
    """给段落或单元格加底纹。"""
    pr = element.get_or_add_tcPr() if element.tag.endswith("}tc") else element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def set_cell_bg(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def add_hyperlink(paragraph, url: str, text: str, size: float) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    # 字体必须在这里直接写进 rPr。超链接的 run 是 w:hyperlink 的子元素，不是 w:p 的
    # 直接子元素，python-docx 的 paragraph.runs 看不到它，事后再补是补不上的。
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_BODY_EN)
    rfonts.set(qn("w:hAnsi"), FONT_BODY_EN)
    rfonts.set(qn("w:eastAsia"), FONT_BODY_CN)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(u)

    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    link.append(new_run)
    paragraph._p.append(link)


# ---------------------------------------------------------------------------
# 行内标记
# ---------------------------------------------------------------------------

INLINE = re.compile(
    r"(?P<img>!\[[^\]]*\]\([^)]+\))"
    r"|(?P<link>\[(?P<ltext>[^\]]+)\]\((?P<lurl>[^)]+)\))"
    r"|(?P<bold>\*\*(?P<btext>.+?)\*\*)"
    r"|(?P<code>`(?P<ctext>[^`]+)`)"
)


def emit_inline(paragraph, text: str, size: float, base_dir: Path,
                bold_all: bool = False) -> None:
    """把一段带行内标记的文本写进段落。"""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_run_font(r, cn=FONT_BODY_CN, en=FONT_BODY_EN, size=size,
                         bold=bold_all or None)
        if m.group("bold"):
            r = paragraph.add_run(m.group("btext"))
            set_run_font(r, cn=FONT_BODY_CN, en=FONT_BODY_EN, size=size, bold=True)
        elif m.group("code"):
            r = paragraph.add_run(m.group("ctext"))
            set_run_font(r, cn=FONT_MONO, en=FONT_MONO, size=size - 0.5, mono=True)
        elif m.group("link"):
            url = m.group("lurl")
            label = m.group("ltext")
            # 本地相对路径不做超链接，转成灰色引用，避免 Word 里点出死链
            if re.match(r"^https?://", url):
                add_hyperlink(paragraph, url, label, size)
            else:
                r = paragraph.add_run(label)
                set_run_font(r, cn=FONT_BODY_CN, en=FONT_BODY_EN, size=size,
                             bold=bold_all or None)
                r2 = paragraph.add_run(f"（{url}）")
                set_run_font(r2, cn=FONT_BODY_CN, en=FONT_BODY_EN, size=size - 1.5,
                             color=MUTED)
        elif m.group("img"):
            pass  # 行内图片在块级处理
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_font(r, cn=FONT_BODY_CN, en=FONT_BODY_EN, size=size,
                     bold=bold_all or None)


# ---------------------------------------------------------------------------
# 块级解析
# ---------------------------------------------------------------------------

RE_H = re.compile(r"^(#{1,6})\s+(.*)$")
RE_IMG = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)\s*$")
RE_FENCE = re.compile(r"^```(\w*)\s*$")
RE_UL = re.compile(r"^(\s*)[-*]\s+(.*)$")
RE_OL = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
RE_HR = re.compile(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$")
RE_QUOTE = re.compile(r"^>\s?(.*)$")
RE_TABLE_SEP = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?\s*$")


def split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def alignments(sep: str) -> list[str]:
    out = []
    for c in split_row(sep):
        left, right = c.startswith(":"), c.endswith(":")
        out.append("center" if left and right else "right" if right else "left")
    return out


class Converter:
    def __init__(self, md_path: Path, doc: Document, base_dir: Path) -> None:
        self.md = md_path
        self.doc = doc
        self.base = base_dir
        self.body_size = 10.5
        self.warnings: list[str] = []

    # -- 各类块 ------------------------------------------------------------

    def heading(self, level: int, text: str) -> None:
        sizes = {1: 20, 2: 15, 3: 12.5, 4: 11.5, 5: 11, 6: 11}
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16 if level <= 2 else 10)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        r = p.add_run(text)
        set_run_font(r, cn=FONT_HEAD_CN, en=FONT_HEAD_EN, size=sizes.get(level, 11),
                     bold=True, color=ACCENT if level <= 3 else None)
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 让 Word 的导航窗格能识别层级
        p.style = self.doc.styles[f"Heading {min(level, 4)}"]
        for run in p.runs:
            set_run_font(run, cn=FONT_HEAD_CN, en=FONT_HEAD_EN,
                         size=sizes.get(level, 11), bold=True,
                         color=ACCENT if level <= 3 else None)

    def para(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.35
        emit_inline(p, text, self.body_size, self.base)

    def bullet(self, text: str, indent: int, ordered: bool, idx: int = 1) -> None:
        p = self.doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        p.paragraph_format.left_indent = Cm(0.75 + 0.6 * indent)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.3
        emit_inline(p, text, self.body_size, self.base)

    def quote(self, lines: list[str]) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        shade(p._p, "F7F7F9")
        emit_inline(p, " ".join(lines), self.body_size - 0.5, self.base)
        for r in p.runs:
            r.font.italic = True

    def hr(self) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "D0D0D0")
        pbdr.append(bottom)
        p._p.get_or_add_pPr().append(pbdr)

    def code(self, lines: list[str], lang: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        shade(p._p, CODE_BG)
        for i, line in enumerate(lines):
            if i:
                br = p.add_run()
                set_run_font(br, cn=FONT_MONO, en=FONT_MONO, size=8.5, mono=True)
                br.add_break()
            r = p.add_run(line)
            set_run_font(r, cn=FONT_MONO, en=FONT_MONO, size=8.5, mono=True)

    def image(self, src: str, alt: str) -> None:
        path = (self.base / src).resolve()
        if not path.exists():
            self.warnings.append(f"图片缺失，已跳过：{src}")
            return
        self.doc.add_picture(str(path), width=Cm(15.5))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if alt:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
            r = cap.add_run(alt)
            set_run_font(r, cn=FONT_BODY_CN, en=FONT_BODY_EN, size=9, color=MUTED)

    def table(self, header: list[str], rows: list[list[str]], aligns: list[str]) -> None:
        ncol = len(header)
        t = self.doc.add_table(rows=1, cols=ncol)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True

        size = 9 if ncol <= 6 else (8 if ncol <= 8 else 7.5)
        for i, cell_text in enumerate(header):
            cell = t.rows[0].cells[i]
            set_cell_bg(cell, TABLE_HEAD_BG)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            txt = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
            r = p.add_run(txt)
            set_run_font(r, cn=FONT_HEAD_CN, en=FONT_HEAD_EN, size=size, bold=True,
                         color=RGBColor(0xFF, 0xFF, 0xFF))

        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for i in range(ncol):
                cell = cells[i]
                if ri % 2 == 1:
                    set_cell_bg(cell, "F7F8FA")
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                a = aligns[i] if i < len(aligns) else "left"
                p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                               "right": WD_ALIGN_PARAGRAPH.RIGHT,
                               "center": WD_ALIGN_PARAGRAPH.CENTER}[a]
                emit_inline(p, row[i] if i < len(row) else "", size, self.base)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -- 主循环 ------------------------------------------------------------

    def run(self) -> list[str]:
        lines = self.md.read_text(encoding="utf-8").split("\n")
        i, n = 0, len(lines)
        in_front_matter = False

        while i < n:
            line = lines[i]

            # YAML front matter（计划文件才有，报告没有，稳妥起见支持）
            if i == 0 and line.strip() == "---":
                in_front_matter = True
                i += 1
                while i < n and lines[i].strip() != "---":
                    i += 1
                i += 1
                continue

            if not line.strip():
                i += 1
                continue

            m = RE_FENCE.match(line)
            if m:
                lang = m.group(1)
                buf = []
                i += 1
                while i < n and not RE_FENCE.match(lines[i]):
                    buf.append(lines[i])
                    i += 1
                i += 1
                png = self.mermaid_png(lang)
                if png is not None:
                    self.image(png, "")
                else:
                    self.code(buf, lang)
                continue

            m = RE_IMG.match(line)
            if m:
                self.image(m.group("src"), m.group("alt"))
                i += 1
                continue

            m = RE_H.match(line)
            if m:
                self.heading(len(m.group(1)), m.group(2).strip())
                i += 1
                continue

            if RE_HR.match(line):
                self.hr()
                i += 1
                continue

            # 表格：当前行是表头且下一行是分隔行
            if "|" in line and i + 1 < n and RE_TABLE_SEP.match(lines[i + 1]):
                header = split_row(line)
                aligns = alignments(lines[i + 1])
                i += 2
                rows = []
                while i < n and "|" in lines[i] and lines[i].strip():
                    rows.append(split_row(lines[i]))
                    i += 1
                self.table(header, rows, aligns)
                continue

            m = RE_QUOTE.match(line)
            if m:
                buf = [m.group(1)]
                i += 1
                while i < n and RE_QUOTE.match(lines[i]):
                    buf.append(RE_QUOTE.match(lines[i]).group(1))
                    i += 1
                self.quote([b for b in buf if b.strip()])
                continue

            m = RE_OL.match(line)
            if m:
                self.bullet(m.group(3), len(m.group(1)) // 2, ordered=True)
                i += 1
                continue

            m = RE_UL.match(line)
            if m:
                self.bullet(m.group(2), len(m.group(1)) // 2, ordered=False)
                i += 1
                continue

            # 普通段落：连续非空行合并
            buf = [line]
            i += 1
            while i < n and lines[i].strip() and not (
                RE_H.match(lines[i]) or RE_FENCE.match(lines[i])
                or RE_HR.match(lines[i]) or RE_UL.match(lines[i])
                or RE_OL.match(lines[i]) or RE_QUOTE.match(lines[i])
                or RE_IMG.match(lines[i]) or "|" in lines[i]
            ):
                buf.append(lines[i])
                i += 1
            self.para(" ".join(x.strip() for x in buf))

        return self.warnings

    def mermaid_png(self, lang: str) -> str | None:
        """mermaid 块若有预渲染 PNG 就用图，否则回落到代码框。"""
        if lang != "mermaid":
            return None
        for cand in ("outputs/fig0_compliance_gate.png",):
            if (self.base / cand).exists():
                return cand
        self.warnings.append("mermaid 图未预渲染，已按代码块输出")
        return None


# ---------------------------------------------------------------------------
# 文档骨架
# ---------------------------------------------------------------------------

def new_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY_EN
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY_CN)
    return doc


def add_footer_page_numbers(doc: Document) -> None:
    for sec in doc.sections:
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        for instr in ("PAGE",):
            fld = OxmlElement("w:fldSimple")
            fld.set(qn("w:instr"), instr)
            run._element.addnext(fld)
        set_run_font(run, cn=FONT_BODY_CN, en=FONT_BODY_EN, size=9, color=MUTED)


def convert(md_path: Path, out_path: Path) -> list[str]:
    doc = new_document()
    conv = Converter(md_path, doc, md_path.parent)
    warnings = conv.run()
    add_footer_page_numbers(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return warnings


def count_markdown_elements(md_path: Path) -> dict:
    """统计 Markdown 里的块级元素，用于和转换结果比对。

    转换器出错时最常见的表现是「悄悄少了一张表或一张图」，而不是报错。
    因此转换完必须回头数一遍，否则无法声称转换是无损的。
    """
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    in_fence = False
    heads = images = tables = 0
    for i, line in enumerate(lines):
        if RE_FENCE.match(line):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        if RE_H.match(line):
            heads += 1
        if RE_IMG.match(line):
            images += 1
        if "|" in line and i + 1 < len(lines) and RE_TABLE_SEP.match(lines[i + 1]):
            tables += 1
    mermaid = len(re.findall(r"^```mermaid", text, re.M))
    return {"headings": heads, "tables": tables,
            "images": images, "mermaid": mermaid}


def count_docx_elements(docx_path: Path) -> dict:
    d = Document(str(docx_path))
    return {
        "headings": sum(1 for p in d.paragraphs if p.style.name.startswith("Heading")),
        "tables": len(d.tables),
        "images": len(d.inline_shapes),
    }


def check(md_path: Path, docx_path: Path) -> tuple[bool, str]:
    m = count_markdown_elements(md_path)
    d = count_docx_elements(docx_path)
    expect_img = m["images"] + m["mermaid"]
    ok = (d["headings"] == m["headings"] and d["tables"] == m["tables"]
          and d["images"] == expect_img)
    detail = (f"标题 {d['headings']}/{m['headings']}  "
              f"表格 {d['tables']}/{m['tables']}  "
              f"插图 {d['images']}/{expect_img}")
    return ok, detail


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("paths", nargs="*", help="要转换的 .md 文件")
    ap.add_argument("--all", action="store_true", help="转换 reports/ 下全部 .md")
    ap.add_argument("--outdir", default=None,
                    help="输出目录；默认写到每个 md 同级的 docx/ 子目录")
    ap.add_argument("--check", action="store_true",
                    help="转换后核对标题/表格/插图数量是否与源文档一致")
    args = ap.parse_args()

    repo = Path(__file__).resolve().parent.parent
    targets: list[Path] = []
    if args.all:
        targets = sorted((repo / "reports").rglob("*.md"))
    targets += [Path(p).resolve() for p in args.paths]
    targets = [t for t in targets if t.exists()]
    if not targets:
        print("没有找到要转换的 .md")
        return 1

    total_warn = 0
    failed = []
    for md in targets:
        out_dir = Path(args.outdir).resolve() if args.outdir else md.parent / "docx"
        out = out_dir / (md.stem + ".docx")
        warns = convert(md, out)
        rel = out.relative_to(repo) if out.is_relative_to(repo) else out
        size_kb = out.stat().st_size / 1024
        line = f"  {md.name}  ->  {rel}  ({size_kb:,.0f} KB)"
        if args.check:
            ok, detail = check(md, out)
            line += f"  [{'OK' if ok else '不一致'}  {detail}]"
            if not ok:
                failed.append(md.name)
        print(line)
        for w in warns:
            print(f"      提示：{w}")
        total_warn += len(warns)

    print(f"\n完成：{len(targets)} 个文件"
          f"{'' if not total_warn else f'，{total_warn} 条提示'}")
    if failed:
        print(f"核对未通过：{', '.join(failed)}")
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
